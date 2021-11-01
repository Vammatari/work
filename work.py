from os import walk
from docx.api import Document
import openpyxl as op

reports = []
for root, dirs, files in walk('C:/Users/Vammatar/Downloads/'):
    for file in files:
        if '.docx' in file:
            reports.append(root + '/' + file)



workbook = op.load_workbook('C:/Users/Vammatar/Downloads/pertorgaphy.xlsx')
worksheet = workbook.active
for report in reports:
    document = Document(report)
    tables = document.tables
    for table in document.tables:
        for row in table.rows:
            text = [cell.text for cell in row.cells]
            worksheet.append(text)
        worksheet.append(['*******************************************','*********','**********','*********'])
workbook.save('C:/Users/Vammatar/Downloads/pertorgaphy.xlsx')